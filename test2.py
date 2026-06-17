from pathlib import Path
from collections import defaultdict

from docling.document_converter import DocumentConverter
from docling_core.types.doc import TableItem

from docx import Document
from docx.shared import Pt


INPUT_PDF = "hindi.pdf"
OUTPUT_DOCX = "output/hindi_parsed.docx"


Path("output").mkdir(exist_ok=True)


converter = DocumentConverter()

result = converter.convert(INPUT_PDF)

doc = result.document


pages = defaultdict(list)


for item, _level in doc.iterate_items():

    if item.prov:
        page_no = item.prov[0].page_no
        pages[page_no].append(item)



word = Document()


# font setup for Hindi
style = word.styles["Normal"]
style.font.name = "Nirmala UI"
style.font.size = Pt(10)



for page_no, items in pages.items():

    # page separator
    word.add_heading(
        f"Page {page_no}",
        level=1
    )


    for item in items:


        # normal text
        if hasattr(item, "text") and item.text:

            word.add_paragraph(
                item.text
            )


        # tables
        elif isinstance(item, TableItem):

            table_md = item.export_to_markdown()


            # simple markdown table parser
            lines = [
                x for x in table_md.split("\n")
                if x.strip()
            ]


            if len(lines) > 1:

                rows = []

                for line in lines:

                    if "---" in line:
                        continue

                    cols = [
                        c.strip()
                        for c in line.strip("|").split("|")
                    ]

                    rows.append(cols)


                if rows:

                    table = word.add_table(
                        rows=len(rows),
                        cols=len(rows[0])
                    )


                    table.style = "Table Grid"


                    for r, row in enumerate(rows):

                        for c, cell in enumerate(row):

                            table.cell(r,c).text = cell



    # page break
    word.add_page_break()



word.save(OUTPUT_DOCX)


print(
    f"Created {OUTPUT_DOCX}"
)
